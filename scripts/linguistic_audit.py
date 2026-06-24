#!/usr/bin/env python3
"""linguistic_audit_fix.py — Complete linguistic audit and fix of resumo_expandido."""

import docx
from docx.shared import Pt
from docx.oxml.ns import qn
import re, os

SRC = 'resumo_expandido_tireoide_2026.docx'
doc = docx.Document(SRC)

def set_p(para, text, fs=Pt(12), b=False):
    """Replace paragraph text, Times New Roman, given size and bold."""
    for r in para.runs: r.text = ''
    if para.runs:
        para.runs[0].text = text
        para.runs[0].font.name = 'Times New Roman'
        para.runs[0].font.size = fs
        para.runs[0].bold = b
    for r in para.runs[1:]:
        r.text = ''

# ═══════════════════════════════════════════════════════════════════════════════
# P6 — ABSTRACT (10pt). Fix accents, crase, flow.
# ═══════════════════════════════════════════════════════════════════════════════
set_p(doc.paragraphs[6],
    'Objetivou-se analisar a expressão diferencial dos genes da via de '
    'sinalização do hormônio tireoidiano (SHT, KEGG hsa04919) no carcinoma '
    'de tireoide, identificando alterações transcriptômicas e padrões de '
    'interação proteína-proteína capazes de gerar hipóteses biológicas sobre '
    'processos moleculares potencialmente associados à biologia tumoral. '
    'Trata-se de estudo exploratório em bioinformática, com dados de RNA-seq '
    'do conjunto integrado TCGA-GTEx (THCA, n = 783) processados no ambiente '
    'R. Foram identificados 29 genes diferencialmente expressos (GDEs) na via '
    'SHT (9 superexpressos, 20 subexpressos; |log2FC| > 1; FDR < 0,05) por '
    'meio de volcano plot. A rede de interação proteína-proteína (PPI) foi '
    'construída com STRING v12.0 (escore ≥ 700) e revelou 8 proteínas '
    'interagentes organizadas em dois módulos funcionais, detectados pelo '
    'algoritmo walktrap. PRKCA destacou-se como gene de maior centralidade '
    'na rede (grau = 5, betweenness = 0,476), atuando como elo entre os '
    'módulos. Os demais 21 GDEs não formaram interações de alta confiança '
    'na rede. Os achados geram hipóteses biológicas sobre a via de sinalização '
    'do hormônio tireoidiano no carcinoma de tireoide, contribuindo para '
    'futuras investigações em oncologia molecular e enfermagem de precisão.',
    Pt(10))

# ═══════════════════════════════════════════════════════════════════════════════
# P10 — OBJECTIVES. Fix accents, crase, improve flow.
# ═══════════════════════════════════════════════════════════════════════════════
set_p(doc.paragraphs[10],
    'Diante desse cenário, objetivou-se analisar a expressão diferencial '
    'dos genes da via de sinalização do hormônio tireoidiano (hsa04919) '
    'no carcinoma de tireoide, identificando alterações transcriptômicas e '
    'padrões de interação proteína-proteína capazes de gerar hipóteses '
    'biológicas sobre processos moleculares potencialmente associados à '
    'biologia tumoral e discutir sua relevância para a enfermagem de precisão. '
    'Especificamente, buscou-se identificar genes diferencialmente expressos '
    'da via SHT no carcinoma de tireoide em comparação ao tecido tireoidiano '
    'normal, construir uma rede de interação proteína-proteína dos GDEs, '
    'identificar genes com elevada centralidade e potencial relevância '
    'biológica na rede e discutir as contribuições dos achados para a '
    'compreensão molecular do carcinoma de tireoide e suas potenciais '
    'aplicações futuras na enfermagem de precisão.',
    Pt(12))

# ═══════════════════════════════════════════════════════════════════════════════
# P18 — PPI RESULTS. Replace "Outrossim" (archaic) + "E, por sua vez" (weak start).
# ═══════════════════════════════════════════════════════════════════════════════
set_p(doc.paragraphs[18],
    'A construção da rede PPI permitiu identificar 24 genes interagentes '
    'dentre os 29 DEGs, organizados em dois módulos biológicos distintos. '
    'Cinco GDEs não participaram da rede por ausência de interações '
    'proteína-proteína acima do limiar de confiança adotado: RXRG '
    '(log2FC = 5,28; FDR = 3,26e-149), DIO1 (log2FC = -2,60; '
    'FDR = 2,44e-27), MED13 (log2FC = 1,20; FDR = 1,43e-76), TBC1D4 '
    '(log2FC = -1,63; FDR = 2,06e-94) e PFKFB2 (log2FC = -1,09; '
    'FDR = 4,48e-40). O módulo 1 foi composto por PRKCA '
    '(log2FC = -1,02; FDR = 5,97e-38), PRKCG (log2FC = -1,05; '
    'FDR = 1,92e-30), PLCG1 (log2FC = -1,43; FDR = 5,77e-141), PLCD3 '
    '(log2FC = 2,38; FDR = 1,19e-95), PLCD4 (log2FC = -1,53; '
    'FDR = 2,89e-71) e PIK3R2 (log2FC = 1,34; FDR = 1,80e-14), '
    'constituindo uma rede densamente conectada e associada a componentes '
    'centrais da transdução de sinais intracelulares. O módulo 2, por sua '
    'vez, foi constituído por ACTG1 (log2FC = 1,41; FDR = 1,72e-119) e '
    'ITGAV (log2FC = 1,03; FDR = 3,52e-55), apresentando menor '
    'conectividade e ligação indireta com o módulo principal.',
    Pt(12))

# ═══════════════════════════════════════════════════════════════════════════════
# P19 — CENTRALITY. Fix accents, improve openings.
# ═══════════════════════════════════════════════════════════════════════════════
set_p(doc.paragraphs[19],
    'No que se refere à topologia da rede, PRKCA destacou-se como o gene '
    'de maior centralidade, apresentando grau 5, betweenness de 0,476, '
    'closeness de 8,52×10⁻⁴ e hub score de 1,000. A conexão entre os dois '
    'módulos foi estabelecida exclusivamente pela interação PRKCA-ACTG1 '
    '(escore STRING = 918), conferindo a PRKCA uma posição topológica '
    'singular como único elo entre a sinalização intracelular mediada por '
    'fosfolipases C e proteínas quinase C (Módulo 1) e a organização '
    'estrutural do citoesqueleto de actina e das interações célula-matriz '
    '(Módulo 2). PLCG1, PLCD4 e PRKCG também apresentaram elevada '
    'conectividade (grau = 4), embora restrita ao Módulo 1, reforçando a '
    'coesão funcional desse agrupamento. Ressalta-se que essas métricas de '
    'centralidade são exploratórias e derivadas de uma rede de 8 nós, '
    'devendo ser interpretadas como geradoras de hipóteses, e não como '
    'evidência de relevância funcional comprovada.',
    Pt(12))

# ═══════════════════════════════════════════════════════════════════════════════
# P20 — OUTSIDE NETWORK. Fix accents.
# ═══════════════════════════════════════════════════════════════════════════════
set_p(doc.paragraphs[20],
    'Em contrapartida, 21 dos 29 GDEs não integraram a rede conectada '
    'final, incluindo genes de alta significância estatística como MYH7, '
    'DIO3, CCND1, RXRG e ATP2A1. Embora esses genes não apresentem '
    'interações de alta confiança (escore ≥ 700) entre si na base STRING, '
    'possuem relevância biológica individual documentada na literatura. Tal '
    'achado sugere que as alterações transcriptômicas na via SHT no carcinoma '
    'de tireoide podem envolver tanto mecanismos coordenados de sinalização '
    'intracelular, representados pela rede PPI, quanto mecanismos '
    'independentes exercidos por genes que atuam por vias distintas ou que '
    'interagem por meio de proteínas não capturadas pelos limiares de '
    'confiança adotados. A investigação desses genes em estudos futuros, '
    'com estratégias complementares de análise de redes e validação '
    'experimental, poderá ampliar a compreensão dos processos moleculares '
    'associados à biologia tumoral tireoidiana.',
    Pt(12))

# ═══════════════════════════════════════════════════════════════════════════════
# P24 — ACKNOWLEDGMENTS heading. Fix accents.
# ═══════════════════════════════════════════════════════════════════════════════
set_p(doc.paragraphs[24], 'Agradecimentos', Pt(12), True)

# ═══════════════════════════════════════════════════════════════════════════════
# P25 — ACKNOWLEDGMENTS text. Fix accents.
# ═══════════════════════════════════════════════════════════════════════════════
set_p(doc.paragraphs[25],
    'Agradecemos à Escola Técnica Estadual João Barcelos Martins e ao '
    'Instituto Federal Fluminense de Educação, Ciência e Tecnologia Campus '
    'Campos Guarus pela promoção dessa troca sinérgica.',
    Pt(12))

doc.save(SRC)
print('[OK] Linguistic audit applied — accents restored, connectors improved, flow polished')
