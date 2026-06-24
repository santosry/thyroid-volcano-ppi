#!/usr/bin/env python3
"""
audit_docx_locked.py — Apply locked versions to resumo_expandido_tireoide_2026.docx

Locked elements:
- Title
- Hypothesis
- General Objective
- Specific Objectives
- Conceptual Conclusion

Formatting: Times New Roman, 12pt body, 10pt abstract, italics for foreign words/packages, no dashes.
"""

import docx
from docx.shared import Pt
import os

SRC = os.path.join(os.path.dirname(__file__), '..', 'resumo_expandido_tireoide_2026.docx')
doc = docx.Document(SRC)

# Helper: replace paragraph with single formatted run
def set_p(para, text, font_size=Pt(12), bold=False, italic=False, font_name='Times New Roman'):
    for r in para.runs:
        r.text = ''
    if para.runs:
        r = para.runs[0]
    else:
        r = para.add_run('')
    r.text = text
    r.font.name = font_name
    r.font.size = font_size
    r.bold = bold
    r.italic = italic

# ═══════════════════════════════════════════════════════════════════════════════
# P0 — TITLE (LOCKED)
# ═══════════════════════════════════════════════════════════════════════════════
set_p(doc.paragraphs[0],
    "Analise Transcriptomica da Via de Sinalizacao do Hormonio Tireoidiano no "
    "Carcinoma de Tireoide e Potenciais Implicacoes para a Enfermagem de Precisao",
    font_size=Pt(12), bold=True)

# ═══════════════════════════════════════════════════════════════════════════════
# P5 — "Resumo" heading
# ═══════════════════════════════════════════════════════════════════════════════
set_p(doc.paragraphs[5], "Resumo", font_size=Pt(11), bold=True)

# ═══════════════════════════════════════════════════════════════════════════════
# P6 — ABSTRACT (10pt, incorporating locked hypothesis + general objective)
# ═══════════════════════════════════════════════════════════════════════════════
set_p(doc.paragraphs[6],
    "Objetivou-se analisar a expressao diferencial dos genes da via de sinalizacao "
    "do hormonio tireoidiano (SHT, KEGG hsa04919) no carcinoma de tireoide, "
    "identificando alteracoes transcriptomicas e padroes de interacao "
    "proteina-proteina capazes de gerar hipoteses biologicas sobre mecanismos "
    "moleculares envolvidos na doenca. Trata-se de estudo exploratorio em "
    "bioinformatica, com dados de RNA-seq do conjunto integrado TCGA-GTEx "
    "(THCA, n = 783) processados no ambiente R. Foram identificados 29 genes "
    "diferencialmente expressos (GDEs) na via SHT (9 superexpressos, 20 "
    "subexpressos; |log2FC| > 1; FDR < 0,05) por meio de volcano plot. A rede "
    "PPI (STRING v12.0, escore >= 700) revelou 8 proteinas interagentes "
    "organizadas em dois modulos funcionais detectados pelo algoritmo walktrap, "
    "com PRKCA como principal gene hub (grau = 5, betweenness = 0,476), atuando "
    "como elo entre os modulos. Os demais 21 GDEs nao formaram interacoes de alta "
    "confianca na rede. O carcinoma de tireoide apresenta alteracoes "
    "transcriptomicas na via SHT, refletindo processos moleculares associados a "
    "biologia tumoral. A identificacao desses genes e de suas interacoes proteicas "
    "contribui para a geracao de hipoteses biologicas relevantes para futuras "
    "investigacoes em oncologia molecular e enfermagem de precisao.",
    font_size=Pt(10), bold=False)

# ═══════════════════════════════════════════════════════════════════════════════
# P7 — KEYWORDS (bold label, regular keywords)
# ═══════════════════════════════════════════════════════════════════════════════
# First run: bold "Palavras-chave: "
# Second run: regular keywords
p7 = doc.paragraphs[7]
for r in p7.runs:
    r.text = ''
if len(p7.runs) >= 2:
    p7.runs[0].text = 'Palavras-chave: '
    p7.runs[0].font.name = 'Times New Roman'
    p7.runs[0].font.size = Pt(11)
    p7.runs[0].bold = True
    p7.runs[1].text = 'Cancer de Tireoide, Doencas Cardiovasculares, Informatica em Enfermagem.'
    p7.runs[1].font.name = 'Times New Roman'
    p7.runs[1].font.size = Pt(11)
    p7.runs[1].bold = False

# ═══════════════════════════════════════════════════════════════════════════════
# P8 — "1. Introducao"
# ═══════════════════════════════════════════════════════════════════════════════
set_p(doc.paragraphs[8], "1. Introducao", font_size=Pt(12), bold=True)

# ═══════════════════════════════════════════════════════════════════════════════
# P9 — INTRODUCTION P1 (locked hypothesis framework)
# ═══════════════════════════════════════════════════════════════════════════════
set_p(doc.paragraphs[9],
    "O carcinoma de tireoide origina-se nas celulas parenquimatosas da tireoide "
    "e apresenta incidencia crescente mundialmente[1]. A meta-analise de Tsai "
    "et al. (2023)[2] reportou maior risco de doenca cerebrovascular "
    "(RR = 1,15; IC95%: 1,10-1,21) e fibrilacao atrial (RR = 1,59; "
    "IC95%: 1,45-1,73) em individuos com cancer de tireoide comparados a "
    "populacao geral, sugerindo associacao epidemiologica entre essas condicoes. "
    "Estudos de bioinformatica integrativa, como o de Zhang et al. (2021)[3], "
    "demonstraram a viabilidade de se identificar genes-chave e vias de sinalizacao "
    "relevantes no carcinoma de tireoide por meio de analise transcriptomica com "
    "dados publicos, evidenciando o potencial dessas abordagens para a descoberta "
    "de alvos moleculares. Entretanto, os mecanismos moleculares subjacentes a "
    "essa associacao permanecem pouco compreendidos. O carcinoma de tireoide "
    "apresenta alteracoes transcriptomicas em genes da via de sinalizacao do "
    "hormonio tireoidiano (hsa04919), refletindo processos moleculares associados "
    "a biologia tumoral. A identificacao desses genes e de suas interacoes "
    "proteicas pode contribuir para a geracao de hipoteses biologicas relevantes "
    "para futuras investigacoes em oncologia molecular e enfermagem de precisao.",
    font_size=Pt(12))

# ═══════════════════════════════════════════════════════════════════════════════
# P10 — INTRODUCTION P2 (general objective + specific objectives)
# ═══════════════════════════════════════════════════════════════════════════════
set_p(doc.paragraphs[10],
    "Diante desse cenario, objetivou-se analisar a expressao diferencial dos "
    "genes da via de sinalizacao do hormonio tireoidiano (hsa04919) no carcinoma "
    "de tireoide, identificando alteracoes transcriptomicas e padroes de interacao "
    "proteina-proteina capazes de gerar hipoteses biologicas sobre mecanismos "
    "moleculares envolvidos na doenca e discutir sua relevancia para a enfermagem "
    "de precisao. Especificamente, buscou-se: (i) identificar genes diferencialmente "
    "expressos da via SHT no carcinoma de tireoide em comparacao ao tecido "
    "tireoidiano normal; (ii) construir uma rede de interacao proteina-proteina "
    "dos GDEs; (iii) identificar genes com elevada centralidade e potencial "
    "relevancia biologica na rede; e (iv) discutir as contribuicoes dos achados "
    "para a compreensao molecular do carcinoma de tireoide e suas potenciais "
    "aplicacoes futuras na enfermagem de precisao.",
    font_size=Pt(12))

# ═══════════════════════════════════════════════════════════════════════════════
# P11 — "2. Materiais e Metodos"
# ═══════════════════════════════════════════════════════════════════════════════
set_p(doc.paragraphs[11], "2. Materiais e Metodos", font_size=Pt(12), bold=True)

# ═══════════════════════════════════════════════════════════════════════════════
# P12 — "2.1. Materiais"
# ═══════════════════════════════════════════════════════════════════════════════
set_p(doc.paragraphs[12], "2.1. Materiais", font_size=Pt(12), bold=False)

# ═══════════════════════════════════════════════════════════════════════════════
# P13 — 2.1 text
# ═══════════════════════════════════════════════════════════════════════════════
set_p(doc.paragraphs[13],
    "Trata-se de estudo exploratorio, baseado em dados secundarios, desenvolvido "
    "no campo da bioinformatica aplicada a oncologia tireoidiana. Os dados de "
    "expressao genica foram obtidos por meio da plataforma UCSC Xena Browser, "
    "utilizando o conjunto integrado TCGA-GTEx (504 amostras tumorais THCA; 279 "
    "amostras normais de tecido tireoidiano), disponibilizado em escala log2. "
    "A analise de expressao diferencial foi conduzida com o pacote limma (modelo "
    "linear com moderacao empirica de Bayes), considerando como criterios de "
    "significancia |log2 Fold Change| > 1 e valor de p ajustado pelo metodo de "
    "Benjamini-Hochberg (FDR) < 0,05. As interacoes proteina-proteina foram "
    "investigadas por meio do banco de dados STRING v12.0 (Homo sapiens, taxon "
    "9606), adotando-se escore combinado minimo de 700 (alta confianca). Os "
    "resultados foram representados por volcano plot e rede PPI.",
    font_size=Pt(12))

# ═══════════════════════════════════════════════════════════════════════════════
# P14 — "2.2. Metodologia"
# ═══════════════════════════════════════════════════════════════════════════════
set_p(doc.paragraphs[14], "2.2. Metodologia", font_size=Pt(12), bold=False)

# ═══════════════════════════════════════════════════════════════════════════════
# P15 — 2.2 text
# ═══════════════════════════════════════════════════════════════════════════════
set_p(doc.paragraphs[15],
    "Os dados transcriptomicos foram obtidos na plataforma UCSC Xena Browser a "
    "partir do conjunto integrado TCGA-GTEx referente ao carcinoma de tireoide "
    "(THCA), disponivel em repositorio publico de livre acesso. A variavel "
    "TCGA_GTEX_main_category foi utilizada para a estratificacao biologica das "
    "amostras (GTEX Thyroid vs. TCGA Thyroid Carcinoma). Os genes da via SHT foram "
    "identificados via KEGG REST API (hsa04919). Todas as etapas de processamento "
    "foram executadas no ambiente R v4.6.0. Empregaram-se os pacotes KEGGREST "
    "(consulta a via metabolica), readr (importacao de dados), limma (expressao "
    "diferencial), dplyr e tibble (manipulacao de dados), httr e jsonlite (consulta "
    "a STRING REST API), igraph e ggraph (construcao e visualizacao da rede PPI), "
    "e ggplot2 e ggrepel (volcano plot). O pipeline completo, os parametros de "
    "analise e os scripts reprodutiveis estao disponiveis publicamente em "
    "repositorio GitHub (https://github.com/santosry/thyroid-volcano-ppi). "
    "Ressalta-se que a comparacao TCGA versus GTEx pode estar sujeita a batch "
    "effects inerentes as diferencas de plataforma, processamento e perfil "
    "demografico entre as coortes, constituindo limitacao metodologica deste estudo.",
    font_size=Pt(12))

# ═══════════════════════════════════════════════════════════════════════════════
# P16 — "3. Resultados e Discussao"
# ═══════════════════════════════════════════════════════════════════════════════
set_p(doc.paragraphs[16], "3. Resultados e Discussao", font_size=Pt(12), bold=True)

# ═══════════════════════════════════════════════════════════════════════════════
# P17 — RESULTS P1 (volcano + PPI data)
# ═══════════════════════════════════════════════════════════════════════════════
set_p(doc.paragraphs[17],
    "Foram analisados 119 genes da via SHT (hsa04919), dos quais 29 (24,4%) "
    "apresentaram expressao diferencial significativa: 9 superexpressos e 20 "
    "subexpressos no tumor em relacao ao tecido normal (|log2FC| > 1; FDR < 0,05). "
    "A Figura 1 apresenta o volcano plot com a distribuicao dos genes, destacando-se "
    "MYH7 (log2FC = -5,59; FDR = 7,3x10-224), DIO3 (log2FC = -4,97; "
    "FDR = 4,4x10-168), MYH6 (log2FC = -4,42; FDR = 1,5x10-157), RXRG "
    "(log2FC = 5,28; FDR = 3,3x10-149) e CCND1 (log2FC = 2,65; "
    "FDR = 1,0x10-214) entre os mais significativos. A rede PPI (Figura 2), "
    "construida com escore STRING >= 700, resultou em 8 proteinas interagentes "
    "das 29 GDEs, organizadas em dois modulos funcionais detectados por walktrap. "
    "Os 21 GDEs restantes, incluindo MYH7 e DIO3, nao apresentaram interacoes de "
    "alta confianca na rede STRING. O Modulo 1 reuniu genes envolvidos em "
    "sinalizacao intracelular mediada por fosfolipases C e proteinas quinase C "
    "(PRKCA, PRKCG, PLCG1, PLCD4), enquanto o Modulo 2 agregou genes relacionados "
    "a organizacao estrutural celular e interacoes celula-matriz (ACTG1, ITGAV). "
    "PLCD3 e PIK3R2 tambem integraram a rede, porem com menor conectividade.",
    font_size=Pt(12))

# ═══════════════════════════════════════════════════════════════════════════════
# P19 — RESULTS P2 (PRKCA hub)
# ═══════════════════════════════════════════════════════════════════════════════
set_p(doc.paragraphs[19],
    "PRKCA destacou-se como principal gene hub da rede (grau = 5; betweenness = "
    "0,476; closeness = 8,52x10-4; hub score = 1,000), exibindo o maior numero "
    "de interacoes e atuando como elo funcional entre os dois modulos (Figura 2). "
    "Os demais hubs identificados foram PLCG1 (grau = 4; betweenness = 0,286), "
    "PLCD4 (grau = 4; betweenness = 0,095) e PRKCG (grau = 4). A posicao topologica "
    "central de PRKCA e compativel com seu papel biologico na regulacao de "
    "proliferacao, diferenciacao e sobrevivencia celular, processos relevantes "
    "para a oncogenese tireoidiana. A conectividade entre os modulos, estabelecida "
    "exclusivamente pela interacao PRKCA-ACTG1 (escore STRING = 918), sugere que "
    "PRKCA pode coordenar alteracoes na arquitetura celular e nas interacoes com "
    "a matriz extracelular via sinalizacao por fosfolipases C e PKC. Contudo, "
    "ressalta-se que a rede PPI do STRING reflete conhecimento acumulado da "
    "literatura e nao demonstra causalidade, ativacao ou inibicao direta. A "
    "identificacao de hubs constitui observacao exploratoria.",
    font_size=Pt(12))

# ═══════════════════════════════════════════════════════════════════════════════
# P20 — DISCUSSION (nursing precision framework)
# ═══════════════════════════════════════════════════════════════════════════════
set_p(doc.paragraphs[20],
    "A enfermagem de precisao, conforme o marco conceitual de Fu et al. (2020)[4], "
    "propoe a utilizacao de dados omicos para personalizar intervencoes nos "
    "dominios de pesquisa, ensino, pratica clinica e politicas de saude. Embora "
    "os achados deste estudo sejam exploratorios e nao possuam aplicabilidade "
    "clinica imediata, a identificacao de genes com funcoes reconhecidas na "
    "biologia tumoral e nas vias de sinalizacao celular fornece subsidios para "
    "a formulacao de hipoteses que, uma vez validadas experimentalmente, poderao "
    "informar protocolos de vigilancia e estratificacao de risco no cuidado "
    "oncologico. Nesse contexto, o enfermeiro ocupa posicao privilegiada para "
    "integrar o conhecimento multiomico a pratica assistencial, considerando as "
    "facetas sociocultural, ambiental e economica do cotidiano do paciente, "
    "conforme preconizado pela saude de precisao[4].",
    font_size=Pt(12))

# ═══════════════════════════════════════════════════════════════════════════════
# P21 — "4. Conclusoes"
# ═══════════════════════════════════════════════════════════════════════════════
set_p(doc.paragraphs[21], "4. Conclusoes", font_size=Pt(12), bold=True)

# ═══════════════════════════════════════════════════════════════════════════════
# P22 — CONCLUSION (LOCKED conceptual conclusion)
# ═══════════════════════════════════════════════════════════════════════════════
set_p(doc.paragraphs[22],
    "Este estudo nao busca demonstrar mecanismos cardiovasculares, biomarcadores "
    "clinicos ou relacoes causais. Seu proposito e caracterizar alteracoes "
    "transcriptomicas na via de sinalizacao do hormonio tireoidiano e identificar "
    "genes relevantes dentro de uma rede de interacao proteina-proteina, gerando "
    "hipoteses biologicas para investigacoes futuras. Foram identificados 29 genes "
    "da via SHT diferencialmente expressos no carcinoma de tireoide, com PRKCA "
    "emergindo como gene hub central na rede PPI (grau = 5, betweenness = 0,476), "
    "conectando modulos de sinalizacao intracelular e organizacao estrutural. "
    "Genes com alta significancia estatistica (MYH6, MYH7, DIO3, CCND1) e genes "
    "com elevada centralidade na rede (PRKCA, PLCG1, PLCD4, PRKCG) constituem "
    "candidatos prioritarios para investigacoes de validacao experimental. "
    "A enfermagem de precisao e apresentada como campo potencial de aplicacao "
    "translacional dos conhecimentos produzidos, e nao como desfecho diretamente "
    "demonstrado pelos dados.",
    font_size=Pt(12))

doc.save(SRC)
print("[OK] All locked versions applied to resumo_expandido_tireoide_2026.docx")
