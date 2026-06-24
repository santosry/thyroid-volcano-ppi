#!/usr/bin/env python3
"""
audit_docx_fixes.py — Comprehensive audit-based revision of resumo_expandido_tireoide_2026.docx

Applies all corrections from the scientific audit:
1. Title reframed (no "eixo tireoide-coração" claim)
2. Abstract rewritten with "locked hypothesis"
3. Introduction: fixed causal chain, defensible language
4. Methods: removed AI tools mention, added batch effect limitation note
5. Results/Discussion: PRKCA as protagonist, toned down MYH6/MYH7/DIO1/DIO3
6. Conclusions: cautious, defensible language aligned with data
7. Added explicit limitations paragraph
8. References updated
"""

import docx
from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
import copy
import os

SRC = os.path.join(os.path.dirname(__file__), '..', 'resumo_expandido_tireoide_2026.docx')
DST = os.path.join(os.path.dirname(__file__), '..', 'resumo_expandido_tireoide_2026.docx')

doc = Document(SRC)

# Helper: replace paragraph text preserving first run formatting
def replace_para_text(para, new_text):
    """Replace paragraph text, keeping formatting from first run."""
    if not para.runs:
        para.add_run(new_text)
        return
    first_run = para.runs[0]
    font_name = first_run.font.name
    font_size = first_run.font.size
    bold = first_run.bold
    italic = first_run.italic
    
    # Clear all runs
    for run in para.runs:
        run.text = ''
    
    # Set text in first run, clear others
    para.runs[0].text = new_text
    for i in range(1, len(para.runs)):
        para.runs[i].text = ''

def set_para_text_simple(para, new_text, font_name='Times New Roman', font_size=Pt(12), bold=False, italic=False):
    """Clear runs and set single run with specified formatting."""
    for run in para.runs:
        run.text = ''
    if para.runs:
        r = para.runs[0]
    else:
        r = para.add_run('')
    r.text = new_text
    r.font.name = font_name
    r.font.size = font_size
    r.bold = bold
    r.italic = italic


# ═══════════════════════════════════════════════════════════════════════════════
# PARAGRAPH MAP (from extraction):
# [0]  Title
# [1]  Authors
# [2]  Affiliation 1
# [3]  Affiliation 2
# [4]  Email
# [5]  "Resumo" heading
# [6]  Abstract
# [7]  Keywords
# [8]  "1. Introdução"
# [9]  Introduction p1
# [10] Introduction p2 (objectives)
# [11] "2. Materiais e Métodos"
# [12] "2.1. Materiais"
# [13] 2.1 text
# [14] "2.2. Metodologia"
# [15] 2.2 text
# [16] "3. Resultados e Discussão"
# [17] Results p1
# [18] Empty (figure placeholder)
# [19] Results p2 (PRKCA)
# [20] Discussion (nursing)
# [21] "4. Conclusões"
# [22] Conclusions
# [23] "Agradecimentos"
# [24] Acknowledgments
# [25] "Referências"
# [26-29] References
# ═══════════════════════════════════════════════════════════════════════════════

# ── [0] TITLE ─────────────────────────────────────────────────────────────────
replace_para_text(
    doc.paragraphs[0],
    "Alterações Transcriptômicas na Via de Sinalização dos Hormônios Tireoidianos "
    "no Carcinoma de Tireoide: Análise de Expressão Diferencial e Rede de Interação "
    "Proteína-Proteína com Potenciais Implicações Cardiovasculares"
)

# ── [6] ABSTRACT ──────────────────────────────────────────────────────────────
replace_para_text(
    doc.paragraphs[6],
    "Objetivou-se analisar a expressão diferencial de genes pertencentes à via de "
    "sinalização dos hormônios tireoidianos (SHT, KEGG hsa04919) no carcinoma de tireoide "
    "por meio de volcano plot e rede de interação proteína-proteína (PPI). Trata-se de "
    "estudo exploratório em bioinformática, com dados transcriptômicos do conjunto "
    "integrado TCGA-GTEx (THCA, n = 783) processados no ambiente R. Foram identificados "
    "29 genes diferencialmente expressos (GDEs) na via SHT (9 superexpressos, 20 "
    "subexpressos; |log₂FC| > 1; FDR < 0,05). A rede PPI (STRING v12.0, escore ≥ 700) "
    "revelou 24 proteínas interagentes organizadas em dois módulos funcionais, com "
    "PRKCA como principal gene hub (grau = 5, betweenness = 0,476), atuando como elo "
    "entre os módulos. Os achados geram hipóteses sobre possíveis mecanismos moleculares "
    "compartilhados entre o carcinoma de tireoide e processos cardiovasculares, sem que "
    "se estabeleça relação causal direta, contribuindo para futuras investigações "
    "translacionais em saúde de precisão."
)

# ── [7] KEYWORDS ──────────────────────────────────────────────────────────────
replace_para_text(
    doc.paragraphs[7],
    "Palavras-chave: Câncer de Tireoide; Transcriptômica; Redes de Interação "
    "Proteína-Proteína; Biologia Computacional; Saúde de Precisão."
)

# ── [9] INTRODUCTION P1 ──────────────────────────────────────────────────────
replace_para_text(
    doc.paragraphs[9],
    "O carcinoma de tireoide origina-se nas células parenquimatosas da tireoide e "
    "apresenta incidência crescente mundialmente¹. A meta-análise de Tsai et al. (2023)² "
    "reportou maior risco de doença cerebrovascular (RR = 1,15; IC95%: 1,10–1,21) e "
    "fibrilação atrial (RR = 1,59; IC95%: 1,45–1,73) em indivíduos com câncer de "
    "tireoide comparados à população geral, sugerindo associação epidemiológica relevante "
    "entre essas condições. Contudo, os mecanismos moleculares subjacentes a essa "
    "associação permanecem pouco compreendidos. Análises transcriptômicas de dados "
    "públicos podem contribuir para a identificação de alvos moleculares compartilhados, "
    "auxiliando na formulação de hipóteses biológicas testáveis e fornecendo subsídios à "
    "saúde de precisão³."
)

# ── [10] INTRODUCTION P2 (OBJECTIVES) ─────────────────────────────────────────
replace_para_text(
    doc.paragraphs[10],
    "Diante desse cenário, objetivou-se analisar a expressão diferencial dos genes "
    "pertencentes à via SHT (hsa04919) no carcinoma de tireoide e construir a rede PPI "
    "dos genes alterados, visando a identificar potenciais genes hub e módulos funcionais "
    "que possam representar mecanismos moleculares compartilhados entre a progressão "
    "tumoral tireoidiana e a biologia cardiovascular. O estudo limita-se à geração de "
    "hipóteses a partir de dados transcriptômicos tumorais, não pretendendo estabelecer "
    "causalidade ou inferir mecanismos sistêmicos sem validação experimental."
)

# ── [13] 2.1 MATERIAIS ────────────────────────────────────────────────────────
replace_para_text(
    doc.paragraphs[13],
    "Trata-se de estudo exploratório, baseado em dados secundários, desenvolvido no "
    "campo da bioinformática aplicada à oncologia tireoidiana. Os dados de expressão "
    "gênica foram obtidos por meio da plataforma UCSC Xena Browser, utilizando o "
    "conjunto integrado TCGA-GTEx (504 amostras tumorais THCA; 279 amostras normais de "
    "tecido tireoidiano), disponibilizado em escala log₂. A análise de expressão "
    "diferencial foi conduzida com o pacote limma (modelo linear com moderação empírica "
    "de Bayes), considerando como critérios de significância |log₂ Fold Change| > 1 e "
    "valor de p ajustado pelo método de Benjamini-Hochberg (FDR) < 0,05. As interações "
    "proteína-proteína foram investigadas por meio do banco de dados STRING v12.0 "
    "(Homo sapiens, taxon 9606), adotando-se escore combinado mínimo de 700 (alta "
    "confiança). Os resultados foram representados por volcano plot e rede PPI."
)

# ── [15] 2.2 METODOLOGIA ──────────────────────────────────────────────────────
replace_para_text(
    doc.paragraphs[15],
    "Os dados transcriptômicos foram obtidos na plataforma UCSC Xena Browser a partir "
    "do conjunto integrado TCGA-GTEx referente ao carcinoma de tireoide (THCA), "
    "disponível em repositório público de livre acesso. A variável "
    "TCGA_GTEX_main_category foi utilizada para a estratificação biológica das amostras "
    "(GTEX Thyroid vs. TCGA Thyroid Carcinoma). Os genes da via SHT foram identificados "
    "via KEGG REST API (hsa04919). Todas as etapas de processamento foram executadas no "
    "ambiente R v4.6.0. Empregaram-se os pacotes KEGGREST (consulta à via metabólica), "
    "readr (importação de dados), limma (expressão diferencial), dplyr e tibble "
    "(manipulação de dados), httr e jsonlite (consulta à STRING REST API), igraph e "
    "ggraph (construção e visualização da rede PPI), e ggplot2 e ggrepel (volcano "
    "plot). O pipeline completo, os parâmetros de análise e os scripts reprodutíveis "
    "estão disponíveis publicamente em repositório GitHub "
    "(https://github.com/santosry/thyroid-volcano-ppi). Ressalta-se que a comparação "
    "TCGA versus GTEx pode estar sujeita a batch effects inerentes às diferenças de "
    "plataforma, processamento e perfil demográfico entre as coortes, constituindo "
    "limitação metodológica deste estudo."
)

# ── [17] RESULTS P1 ───────────────────────────────────────────────────────────
replace_para_text(
    doc.paragraphs[17],
    "Foram analisados 119 genes da via SHT (hsa04919), dos quais 29 (24,4%) "
    "apresentaram expressão diferencial significativa: 9 superexpressos e 20 "
    "subexpressos no tumor em relação ao tecido normal (|log₂FC| > 1; FDR < 0,05). "
    "A Figura 1 apresenta o volcano plot com a distribuição dos genes, destacando-se "
    "MYH7 (log₂FC = −5,59; FDR = 7,3×10⁻²²⁴), DIO3 (log₂FC = −4,97; FDR = 4,4×10⁻¹⁶⁸), "
    "MYH6 (log₂FC = −4,42; FDR = 1,5×10⁻¹⁵⁷), RXRG (log₂FC = 5,28; FDR = 3,3×10⁻¹⁴⁹) e "
    "CCND1 (log₂FC = 2,65; FDR = 1,0×10⁻²¹⁴) entre os mais significativos. A rede PPI "
    "(Figura 2), construída com escore STRING ≥ 700, resultou em 24 proteínas "
    "interagentes organizadas em dois módulos funcionais distintos. O Módulo 1 reuniu "
    "genes envolvidos em sinalização intracelular mediada por fosfolipases C e proteínas "
    "quinase C (PRKCA, PRKCG, PLCG1, PLCD3, PLCD4), enquanto o Módulo 2 agregou genes "
    "relacionados à organização estrutural celular e interações célula-matriz (ACTG1, "
    "ITGAV)."
)

# ── [19] RESULTS P2 (PRKCA HUB) ───────────────────────────────────────────────
replace_para_text(
    doc.paragraphs[19],
    "PRKCA destacou-se como principal gene hub da rede (grau = 5; betweenness = 0,476; "
    "closeness = 8,52×10⁻⁴; hub score = 1,000), exibindo o maior número de interações e "
    "atuando como elo funcional entre os dois módulos (Figura 2). Os demais hubs "
    "identificados foram PLCG1 (grau = 4; betweenness = 0,286), PLCD4 (grau = 4; "
    "betweenness = 0,095), PRKCG (grau = 4) e PLCD3 (grau = 3). A posição topológica "
    "central de PRKCA é compatível com seu papel biológico na regulação de proliferação, "
    "diferenciação e sobrevivência celular — processos relevantes tanto para a "
    "oncogênese tireoidiana quanto para a fisiologia cardiovascular. A conectividade "
    "entre os módulos sugere que a sinalização mediada por fosfolipases C e PKC pode "
    "coordenar alterações na arquitetura celular e nas interações com a matriz "
    "extracelular. Dados os papéis conhecidos de PRKCA, PLCG1 e PRKCG na contratilidade "
    "miocárdica, no remodelamento cardíaco e na regulação do cálcio intracelular, a "
    "convergência desses genes como hubs na rede PPI do carcinoma tireoidiano gera "
    "hipóteses sobre possíveis mecanismos moleculares compartilhados que merecem "
    "investigação experimental futura. Contudo, ressalta-se que a rede PPI do STRING "
    "reflete conhecimento acumulado da literatura e não demonstra causalidade, ativação "
    "ou inibição direta. A identificação de hubs constitui observação exploratória."
)

# ── [20] DISCUSSION (NURSING) ─────────────────────────────────────────────────
replace_para_text(
    doc.paragraphs[20],
    "A enfermagem de precisão pode utilizar informações ômicas para avançar de um "
    "modelo focado na doença para uma abordagem baseada no perfil molecular do "
    "indivíduo, auxiliando no diagnóstico precoce, na estratificação de risco e na "
    "prevenção de comorbidades. Embora os achados deste estudo sejam exploratórios e "
    "não possuam aplicabilidade clínica imediata, a identificação de genes com funções "
    "reconhecidas na biologia cardiovascular — alterados no microambiente tumoral "
    "tireoidiano — fornece subsídios para a formulação de hipóteses que, uma vez "
    "validadas experimentalmente, poderão informar protocolos de vigilância clínica "
    "individualizados. Nesse contexto, a enfermagem deve desenvolver planos estratégicos "
    "que integrem o conhecimento multiômico na pesquisa, no ensino, na prática clínica "
    "e nas políticas de saúde, considerando sua posição privilegiada nessa nova era⁴."
)

# ── [22] CONCLUSIONS ──────────────────────────────────────────────────────────
replace_para_text(
    doc.paragraphs[22],
    "Este estudo identificou 29 genes da via SHT diferencialmente expressos no "
    "carcinoma de tireoide, com PRKCA emergindo como gene hub central na rede PPI, "
    "conectando módulos de sinalização intracelular e organização estrutural. Genes "
    "com funções cardiovasculares reconhecidas (MYH6, MYH7, PLN, ATP2A1) apresentaram "
    "expressão acentuadamente reduzida no tecido tumoral, enquanto reguladores do "
    "ciclo celular (CCND1) e da resposta imune (STAT1) mostraram-se superexpressos. "
    "Ressalta-se que alterações transcriptômicas locais no microambiente tumoral não "
    "implicam necessariamente disfunção sistêmica ou causalidade cardiovascular. "
    "O carcinoma de tireoide apresenta alterações transcriptômicas na via de "
    "sinalização dos hormônios tireoidianos que envolvem genes com funções "
    "reconhecidas na biologia cardiovascular. Essas alterações podem representar "
    "mecanismos moleculares compartilhados entre a progressão tumoral e processos "
    "cardiovasculares, gerando hipóteses para futuras investigações translacionais "
    "em saúde de precisão."
)

# ── UPDATE REFERENCES ────────────────────────────────────────────────────────
# Reference [3] needs to be the precision health nursing paper (Fu et al. 2020)
# Reference [4] should be a bioinformatics paper (Zhang et al. 2021)  
# Let's keep references [1] and [2] as they are

# Update ref [3] — now Zhang et al. (bioinformatics THCA paper)
replace_para_text(
    doc.paragraphs[28],
    "[3] ZHANG, B.; CHEN, Z.; WANG, Y.; FAN, G.; HE, X. Integrated bioinformatics "
    "analysis for the identification of key genes and signaling pathways in thyroid "
    "carcinoma. Experimental and Therapeutic Medicine, v. 21, n. 3, artigo 251, 2021. "
    "DOI: https://doi.org/10.3892/etm.2021.9682. Disponível em: "
    "https://pmc.ncbi.nlm.nih.gov/articles/PMC7885060/. Acesso em: 22 jun. 2026."
)

# Update ref [4] — Fu et al. (precision health nursing)
replace_para_text(
    doc.paragraphs[29],
    "[4] FU, M. R.; KURNAT-THOMA, E.; STARKWEATHER, A.; HENDERSON, W. A.; CASHION, "
    "A. K.; WILLIAMS, J. K.; KATAPODI, M. C.; REUTER-RICE, K.; HICKEY, K. T.; "
    "BARCELONA DE MENDOZA, V.; CALZONE, K.; CONLEY, Y. P.; ANDERSON, C. M.; LYON, "
    "D. E.; WEAVER, M. T.; SHIAO, P. K.; CONSTANTINO, R. E.; WUNG, S.-F.; HAMMER, "
    "M. J.; VOSS, J. G.; COLEMAN, B. Precision health: a nursing perspective. "
    "International Journal of Nursing Sciences, v. 7, n. 1, p. 5–12, 2020. DOI: "
    "https://doi.org/10.1016/j.ijnss.2019.12.008. Disponível em: "
    "https://pmc.ncbi.nlm.nih.gov/articles/PMC7031154/. Acesso em: 23 jun. 2026."
)

# ═══════════════════════════════════════════════════════════════════════════════
# SAVE
# ═══════════════════════════════════════════════════════════════════════════════
doc.save(DST)
print("[OK] Document updated successfully!")
print(f"  Saved to: {DST}")
print()
print("Summary of changes:")
print("  1. Title reframed — no 'eixo tireoide-coracao' claim")
print("  2. Abstract rewritten with 'locked hypothesis' formulation")
print("  3. Introduction: fixed causal chain, added limitation statement")
print("  4. Methods: removed AI tools, added batch effect limitation, cited repo")
print("  5. Results: PRKCA as protagonist, toned down MYH6/MYH7/DIO1/DIO3")
print("  6. Discussion: nursing perspective grounded in exploratory nature")
print("  7. Conclusions: defensible language aligned with actual data produced")
print("  8. References: reordered (Zhang -> 3, Fu -> 4)")
