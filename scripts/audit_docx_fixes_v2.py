#!/usr/bin/env python3
"""
audit_docx_fixes_v2.py — Strict audit corrections for resumo_expandido_tireoide_2026.docx

Corrections based on cross-referencing claims against actual data files:
1. "24 proteínas interagentes" → "8 proteínas" (N04_network_summary.tsv confirms 8 nodes)
2. PLCD3 removed from hub list (is_hub=FALSE in N03_centrality_metrics.tsv)
3. Added note: 21 DEGs excluded from PPI network
4. Minor rounding consistency fixes
"""

import docx
from docx.shared import Pt
import os

SRC = os.path.join(os.path.dirname(__file__), '..', 'resumo_expandido_tireoide_2026.docx')
doc = docx.Document(SRC)

def replace_para(para_idx, new_text):
    """Replace all text in paragraph, keeping first run formatting."""
    para = doc.paragraphs[para_idx]
    if not para.runs:
        para.add_run(new_text)
        return
    # Clear all runs, set text in first
    for run in para.runs:
        run.text = ''
    para.runs[0].text = new_text


# ═══════════════════════════════════════════════════════════════════════════════
# FIX 1: Abstract — "24 proteínas interagentes" → "8 proteínas interagentes"
# Data source: N04_network_summary.tsv → Nodes: 8
# ═══════════════════════════════════════════════════════════════════════════════
replace_para(6,
    "Objetivou-se analisar a expressão diferencial de genes pertencentes à via de "
    "sinalização dos hormônios tireoidianos (SHT, KEGG hsa04919) no carcinoma de tireoide "
    "por meio de volcano plot e rede de interação proteína-proteína (PPI). Trata-se de "
    "estudo exploratório em bioinformática, com dados transcriptômicos do conjunto "
    "integrado TCGA-GTEx (THCA, n = 783) processados no ambiente R. Foram identificados "
    "29 genes diferencialmente expressos (GDEs) na via SHT (9 superexpressos, 20 "
    "subexpressos; |log₂FC| > 1; FDR < 0,05). A rede PPI (STRING v12.0, escore ≥ 700) "
    "revelou 8 proteínas interagentes organizadas em dois módulos funcionais, com "
    "PRKCA como principal gene hub (grau = 5, betweenness = 0,476), atuando como elo "
    "entre os módulos. Os demais 21 GDEs não formaram interações de alta confiança na "
    "rede. Os achados geram hipóteses sobre possíveis mecanismos moleculares "
    "compartilhados entre o carcinoma de tireoide e processos cardiovasculares, sem que "
    "se estabeleça relação causal direta, contribuindo para futuras investigações "
    "translacionais em saúde de precisão."
)

# ═══════════════════════════════════════════════════════════════════════════════
# FIX 2: Results — "24 proteínas interagentes" → "8 proteínas interagentes"
# + add note about 21 DEGs excluded
# + PLCD3 NOT listed as hub (is_hub=FALSE in N03, degree=3, betweenness=0)
# ═══════════════════════════════════════════════════════════════════════════════
replace_para(17,
    "Foram analisados 119 genes da via SHT (hsa04919), dos quais 29 (24,4%) "
    "apresentaram expressão diferencial significativa: 9 superexpressos e 20 "
    "subexpressos no tumor em relação ao tecido normal (|log₂FC| > 1; FDR < 0,05). "
    "A Figura 1 apresenta o volcano plot com a distribuição dos genes, destacando-se "
    "MYH7 (log₂FC = −5,59; FDR = 7,3×10⁻²²⁴), DIO3 (log₂FC = −4,97; FDR = 4,4×10⁻¹⁶⁸), "
    "MYH6 (log₂FC = −4,42; FDR = 1,5×10⁻¹⁵⁷), RXRG (log₂FC = 5,28; FDR = 3,3×10⁻¹⁴⁹) e "
    "CCND1 (log₂FC = 2,65; FDR = 1,0×10⁻²¹⁴) entre os mais significativos. A rede PPI "
    "(Figura 2), construída com escore STRING ≥ 700, resultou em 8 proteínas "
    "interagentes — das 29 GDEs — organizadas em dois módulos funcionais. Os 21 GDEs "
    "restantes, incluindo MYH7 e DIO3, não apresentaram interações de alta confiança "
    "na rede STRING. O Módulo 1 reuniu genes envolvidos em sinalização intracelular "
    "mediada por fosfolipases C e proteínas quinase C (PRKCA, PRKCG, PLCG1, PLCD4), "
    "enquanto o Módulo 2 agregou genes relacionados à organização estrutural celular "
    "e interações célula-matriz (ACTG1, ITGAV). PLCD3 e PIK3R2 também integraram a "
    "rede, porém com menor conectividade."
)

# ═══════════════════════════════════════════════════════════════════════════════
# FIX 3: PRKCA paragraph — PLCD3 removed from hub list
# PLCD3: degree=3, betweenness=0, is_hub=FALSE (N03_centrality_metrics.tsv)
# ═══════════════════════════════════════════════════════════════════════════════
replace_para(19,
    "PRKCA destacou-se como principal gene hub da rede (grau = 5; betweenness = 0,476; "
    "closeness = 8,52×10⁻⁴; hub score = 1,000), exibindo o maior número de interações e "
    "atuando como elo funcional entre os dois módulos (Figura 2). Os demais hubs "
    "identificados foram PLCG1 (grau = 4; betweenness = 0,286), PLCD4 (grau = 4; "
    "betweenness = 0,095) e PRKCG (grau = 4). A posição topológica central de PRKCA é "
    "compatível com seu papel biológico na regulação de proliferação, diferenciação e "
    "sobrevivência celular — processos relevantes tanto para a oncogênese tireoidiana "
    "quanto para a fisiologia cardiovascular. A conectividade entre os módulos, "
    "estabelecida exclusivamente pela interação PRKCA-ACTG1 (escore STRING = 918), "
    "sugere que PRKCA pode coordenar alterações na arquitetura celular e nas interações "
    "com a matriz extracelular via sinalização por fosfolipases C e PKC. Dados os papéis "
    "conhecidos de PRKCA, PLCG1 e PRKCG na contratilidade miocárdica, no remodelamento "
    "cardíaco e na regulação do cálcio intracelular, a convergência desses genes como "
    "hubs na rede PPI do carcinoma tireoidiano gera hipóteses sobre possíveis "
    "mecanismos moleculares compartilhados que merecem investigação experimental futura. "
    "Contudo, ressalta-se que a rede PPI do STRING reflete conhecimento acumulado da "
    "literatura e não demonstra causalidade, ativação ou inibição direta. A "
    "identificação de hubs constitui observação exploratória."
)

doc.save(SRC)
print("[OK] Strict audit corrections applied to resumo_expandido_tireoide_2026.docx")
print()
print("Corrections:")
print("  1. Abstract: '24 proteinas' -> '8 proteinas' (N04 confirms 8 nodes)")
print("  2. Abstract: Added note about 21 DEGs excluded from network")
print("  3. Results P17: '24 proteinas' -> '8 proteinas' + 21 excluded note")
print("  4. Results P17: PLCD3 removed from Module 1 hub list")
print("  5. PRKCA P19: PLCD3 removed from hub list (is_hub=FALSE)")
print("  6. PRKCA P19: Added detail about single inter-module edge PRKCA-ACTG1")
