#!/usr/bin/env python3
"""lock_docx_final.py — Apply locked versions with accents, no lists, no CV overclaims."""
import docx
from docx.shared import Pt

doc = docx.Document('resumo_expandido_tireoide_2026.docx')

def S(para, text, fs=Pt(12), b=False):
    for r in para.runs: r.text = ''
    if para.runs:
        para.runs[0].text = text
        para.runs[0].font.name = 'Times New Roman'
        para.runs[0].font.size = fs
        para.runs[0].bold = b

# P0 - TITLE
S(doc.paragraphs[0],
  'Análise Transcriptômica da Via de Sinalização do Hormônio Tireoidiano no '
  'Carcinoma de Tireoide e Potenciais Implicações para a Enfermagem de Precisão',
  Pt(12), True)

# P5 - Resumo
S(doc.paragraphs[5], 'Resumo', Pt(11), True)

# P6 - ABSTRACT (10pt, locked, no CV)
S(doc.paragraphs[6],
  'Objetivou-se analisar a expressão diferencial dos genes da via de sinalização '
  'do hormônio tireoidiano (SHT, KEGG hsa04919) no carcinoma de tireoide, '
  'identificando alterações transcriptômicas e padrões de interação '
  'proteína-proteína capazes de gerar hipóteses biológicas sobre mecanismos '
  'moleculares envolvidos na doença. Trata-se de estudo exploratório em '
  'bioinformática, com dados de RNA-seq do conjunto integrado TCGA-GTEx '
  '(THCA, n = 783) processados no ambiente R. Foram identificados 29 genes '
  'diferencialmente expressos (GDEs) na via SHT (9 superexpressos, 20 '
  'subexpressos; |log2FC| > 1; FDR < 0,05) por meio de volcano plot. A rede '
  'de interação proteína-proteína (PPI) foi construída com STRING v12.0 '
  '(escore >= 700), revelando 8 proteínas interagentes organizadas em dois '
  'módulos funcionais detectados pelo algoritmo walktrap. PRKCA destacou-se '
  'como gene de maior centralidade na rede (grau = 5, betweenness = 0,476), '
  'atuando como elo entre os módulos. Os demais 21 GDEs não formaram interações '
  'de alta confiança na rede. Os achados geram hipóteses biológicas sobre a '
  'via de sinalização do hormônio tireoidiano no carcinoma de tireoide, '
  'contribuindo para futuras investigações em oncologia molecular e enfermagem '
  'de precisão.',
  Pt(10))

# P7 - KEYWORDS
p7 = doc.paragraphs[7]
for r in p7.runs: r.text = ''
if len(p7.runs) >= 2:
    p7.runs[0].text = 'Palavras-chave: '
    p7.runs[0].font.name = 'Times New Roman'; p7.runs[0].font.size = Pt(11); p7.runs[0].bold = True
    p7.runs[1].text = 'Câncer de Tireoide, Doenças Cardiovasculares, Informática em Enfermagem.'
    p7.runs[1].font.name = 'Times New Roman'; p7.runs[1].font.size = Pt(11); p7.runs[1].bold = False

# P8
S(doc.paragraphs[8], '1. Introdução', Pt(12), True)

# P9 - INTRO (locked hypothesis, accents, no CV overreach)
S(doc.paragraphs[9],
  'O carcinoma de tireoide origina-se nas células parenquimatosas da tireoide '
  'e apresenta incidência crescente mundialmente[1]. A meta-análise de Tsai '
  'et al. (2023)[2] reportou maior risco de doença cerebrovascular '
  '(RR = 1,15; IC95%: 1,10-1,21) e fibrilação atrial (RR = 1,59; '
  'IC95%: 1,45-1,73) em indivíduos com câncer de tireoide comparados à '
  'população geral, sugerindo associação epidemiológica entre essas condições. '
  'Estudos de bioinformática integrativa, como o de Zhang et al. (2021)[3], '
  'demonstraram a viabilidade de se identificar genes-chave e vias de sinalização '
  'relevantes no carcinoma de tireoide por meio de análise transcriptômica com '
  'dados públicos. O carcinoma de tireoide apresenta alterações transcriptômicas '
  'em genes da via de sinalização do hormônio tireoidiano (hsa04919), refletindo '
  'processos moleculares associados à biologia tumoral. A identificação desses '
  'genes e de suas interações proteicas pode contribuir para a geração de '
  'hipóteses biológicas relevantes para futuras investigações em oncologia '
  'molecular e enfermagem de precisão.',
  Pt(12))

# P10 - OBJECTIVES (continuous text, no list, accents)
S(doc.paragraphs[10],
  'Diante desse cenário, objetivou-se analisar a expressão diferencial dos '
  'genes da via de sinalização do hormônio tireoidiano (hsa04919) no carcinoma '
  'de tireoide, identificando alterações transcriptômicas e padrões de interação '
  'proteína-proteína capazes de gerar hipóteses biológicas sobre mecanismos '
  'moleculares envolvidos na doença e discutir sua relevância para a enfermagem '
  'de precisão. Especificamente, buscou-se identificar genes diferencialmente '
  'expressos da via SHT no carcinoma de tireoide em comparação ao tecido '
  'tireoidiano normal, construir uma rede de interação proteína-proteína dos '
  'GDEs, identificar genes com elevada centralidade e potencial relevância '
  'biológica na rede e discutir as contribuições dos achados para a compreensão '
  'molecular do carcinoma de tireoide e suas potenciais aplicações futuras na '
  'enfermagem de precisão.',
  Pt(12))

# P11
S(doc.paragraphs[11], '2. Materiais e Métodos', Pt(12), True)

# P12
S(doc.paragraphs[12], '2.1. Materiais', Pt(12))

# P13 - 2.1 (no CV, with accents)
S(doc.paragraphs[13],
  'Trata-se de estudo exploratório, baseado em dados secundários, desenvolvido '
  'no campo da bioinformática aplicada à oncologia tireoidiana. Os dados de '
  'expressão gênica foram obtidos por meio da plataforma UCSC Xena Browser, '
  'utilizando o conjunto integrado TCGA-GTEx (504 amostras tumorais THCA; 279 '
  'amostras normais de tecido tireoidiano), disponibilizado em escala log2. '
  'A análise de expressão diferencial foi conduzida com o pacote limma (modelo '
  'linear com moderação empírica de Bayes), considerando como critérios de '
  'significância |log2 Fold Change| > 1 e valor de p ajustado pelo método de '
  'Benjamini-Hochberg (FDR) < 0,05. As interações proteína-proteína foram '
  'investigadas por meio do banco de dados STRING v12.0 (Homo sapiens, taxon '
  '9606), adotando-se escore combinado mínimo de 700 (alta confiança). Os '
  'resultados foram representados por volcano plot e rede PPI.',
  Pt(12))

# P14
S(doc.paragraphs[14], '2.2. Metodologia', Pt(12))

# P15 - 2.2 (no CV, with accents)
S(doc.paragraphs[15],
  'Os dados transcriptômicos foram obtidos na plataforma UCSC Xena Browser a '
  'partir do conjunto integrado TCGA-GTEx referente ao carcinoma de tireoide '
  '(THCA), disponível em repositório público de livre acesso. A variável '
  'TCGA_GTEX_main_category foi utilizada para a estratificação biológica das '
  'amostras (GTEX Thyroid vs. TCGA Thyroid Carcinoma). Os genes da via SHT foram '
  'identificados via KEGG REST API (hsa04919). Todas as etapas de processamento '
  'foram executadas no ambiente R v4.6.0. Empregaram-se os pacotes KEGGREST '
  '(consulta à via metabólica), readr (importação de dados), limma (expressão '
  'diferencial), dplyr e tibble (manipulação de dados), httr e jsonlite (consulta '
  'à STRING REST API), igraph e ggraph (construção e visualização da rede PPI), '
  'e ggplot2 e ggrepel (volcano plot). O pipeline completo, os parâmetros de '
  'análise e os scripts reprodutíveis estão disponíveis publicamente em '
  'repositório GitHub (https://github.com/santosry/thyroid-volcano-ppi). '
  'Ressalta-se que a comparação TCGA versus GTEx pode estar sujeita a batch '
  'effects inerentes às diferenças de plataforma, processamento e perfil '
  'demográfico entre as coortes, constituindo limitação metodológica deste estudo.',
  Pt(12))

# P16
S(doc.paragraphs[16], '3. Resultados e Discussão', Pt(12), True)

# P17 - RESULTS (volcano + PPI, accents, no CV)
S(doc.paragraphs[17],
  'Foram analisados 119 genes da via SHT (hsa04919), dos quais 29 (24,4%) '
  'apresentaram expressão diferencial significativa: 9 superexpressos e 20 '
  'subexpressos no tumor em relação ao tecido normal (|log2FC| > 1; FDR < 0,05). '
  'A Figura 1 apresenta o volcano plot com a distribuição dos genes, destacando-se '
  'MYH7 (log2FC = -5,59; FDR = 7,3x10-224), DIO3 (log2FC = -4,97; '
  'FDR = 4,4x10-168), MYH6 (log2FC = -4,42; FDR = 1,5x10-157), RXRG '
  '(log2FC = 5,28; FDR = 3,3x10-149) e CCND1 (log2FC = 2,65; '
  'FDR = 1,0x10-214) entre os mais significativos. A rede PPI (Figura 2), '
  'construída com escore STRING >= 700, resultou em 8 proteínas interagentes '
  'das 29 GDEs, organizadas em dois módulos funcionais detectados por walktrap. '
  'Os 21 GDEs restantes, incluindo MYH7 e DIO3, não apresentaram interações de '
  'alta confiança na rede STRING. O Módulo 1 reuniu genes envolvidos em '
  'sinalização intracelular mediada por fosfolipases C e proteínas quinase C '
  '(PRKCA, PRKCG, PLCG1, PLCD4), enquanto o Módulo 2 agregou genes relacionados '
  'à organização estrutural celular e interações célula-matriz (ACTG1, ITGAV). '
  'PLCD3 e PIK3R2 também integraram a rede, porém com menor conectividade.',
  Pt(12))

# P19 - PRKCA (cautious, no CV overclaim, accents)
S(doc.paragraphs[19],
  'PRKCA destacou-se como gene de maior centralidade na rede (grau = 5; '
  'betweenness = 0,476; closeness = 8,52x10-4; hub score = 1,000), exibindo '
  'o maior número de interações e atuando como elo funcional entre os dois '
  'módulos (Figura 2). Os demais genes com elevada centralidade foram PLCG1 '
  '(grau = 4; betweenness = 0,286), PLCD4 (grau = 4; betweenness = 0,095) e '
  'PRKCG (grau = 4). A posição topológica central de PRKCA é compatível com '
  'seu papel biológico na regulação de proliferação, diferenciação e '
  'sobrevivência celular, processos relevantes para a oncogênese tireoidiana. '
  'A conectividade entre os módulos, estabelecida exclusivamente pela interação '
  'PRKCA-ACTG1 (escore STRING = 918), sugere que PRKCA pode coordenar alterações '
  'na arquitetura celular e nas interações com a matriz extracelular via '
  'sinalização por fosfolipases C e PKC. Contudo, ressalta-se que a rede PPI '
  'do STRING reflete conhecimento acumulado da literatura e não demonstra '
  'causalidade, ativação ou inibição direta. A identificação de genes com '
  'elevada centralidade constitui observação exploratória que requer validação '
  'experimental futura.',
  Pt(12))

# P20 - DISCUSSION (nursing only, no CV, accents)
S(doc.paragraphs[20],
  'A enfermagem de precisão, conforme o marco conceitual de Fu et al. (2020)[4], '
  'propõe a utilização de dados ômicos para personalizar intervenções nos '
  'domínios de pesquisa, ensino, prática clínica e políticas de saúde. Embora '
  'os achados deste estudo sejam exploratórios e não possuam aplicabilidade '
  'clínica imediata, a identificação de genes com elevada centralidade na via '
  'de sinalização do hormônio tireoidiano e de padrões de interação '
  'proteína-proteína no carcinoma de tireoide fornece subsídios para a '
  'formulação de hipóteses biológicas que, uma vez validadas '
  'experimentalmente, poderão informar futuras abordagens translacionais. '
  'Nesse contexto, o enfermeiro ocupa posição privilegiada para integrar o '
  'conhecimento multiômico à prática assistencial, considerando as facetas '
  'sociocultural, ambiental e econômica do cotidiano do paciente, conforme '
  'preconizado pela saúde de precisão[4].',
  Pt(12))

# P21
S(doc.paragraphs[21], '4. Conclusões', Pt(12), True)

# P22 - CONCLUSION (locked conceptual, no CV, accents)
S(doc.paragraphs[22],
  'Este estudo não busca demonstrar mecanismos cardiovasculares, biomarcadores '
  'clínicos ou relações causais. Seu propósito é caracterizar alterações '
  'transcriptômicas na via de sinalização do hormônio tireoidiano e identificar '
  'genes relevantes dentro de uma rede de interação proteína-proteína, gerando '
  'hipóteses biológicas para investigações futuras. Foram identificados 29 genes '
  'da via SHT diferencialmente expressos no carcinoma de tireoide, com PRKCA '
  'apresentando a maior centralidade na rede PPI (grau = 5, betweenness = 0,476), '
  'conectando módulos de sinalização intracelular e organização estrutural. '
  'Genes com alta significância estatística (MYH6, MYH7, DIO3, CCND1) e genes '
  'com elevada centralidade na rede (PRKCA, PLCG1, PLCD4, PRKCG) constituem '
  'candidatos prioritários para investigações de validação experimental. '
  'A enfermagem de precisão é apresentada como campo potencial de aplicação '
  'translacional dos conhecimentos produzidos, e não como desfecho diretamente '
  'demonstrado pelos dados.',
  Pt(12))

doc.save('resumo_expandido_tireoide_2026.docx')
print('[OK] Lock final applied — accents, no lists, no CV overclaims')
